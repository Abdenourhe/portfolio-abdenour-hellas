from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('LISTE DES CONTACTS LINKEDIN À SOLLICITER', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Abdenour Hellas - Génie Électrique | Canada')
run.font.size = Pt(11)
run.font.italic = True

doc.add_paragraph()

doc.add_heading('1. COMMENT CONTACTER LES RECRUTEURS SUR LINKEDIN', level=1)

doc.add_paragraph('Ne postulez pas seulement aux offres. Contactez directement les recruteurs et les gestionnaires d\'embauche.', style='List Bullet')
doc.add_paragraph('LinkedIn permet d\'envoyer des messages même sans connexion (InMail gratuit avec Premium, ou message de connexion).', style='List Bullet')
doc.add_paragraph('Un message personnalisé a 3× plus de chances d\'être lu qu\'une candidature anonyme.', style='List Bullet')

doc.add_paragraph()

# Message modèle
doc.add_heading('2. MESSAGE MODÈLE POUR LINKEDIN', level=1)

message = """Bonjour [Prénom],

Je me permets de vous contacter car je suis à la recherche d'opportunités en génie électrique au Canada.

Je suis récemment diplômé d'un M.Eng. en Génie Électrique de l'UQAT (2025), avec 5+ ans d'expérience en maintenance industrielle, conception de systèmes électriques et automatisation. J'ai travaillé sur des projets variés en Algérie et au Québec, notamment le système collaboratif électrique TCOST et la supervision des opérations électroménagers chez RONA.

Je maîtrise les normes CQE et CEC, ainsi que les logiciels AutoCAD, Proteus 8, MATLAB et SolidWorks. Je développe également des solutions web (Next.js, React) pour l'automatisation des processus techniques.

J'ai postulé à plusieurs postes chez [Entreprise] et je serais ravi d'échanger avec vous sur les opportunités actuelles ou à venir au sein de votre équipe.

Cordialement,
Abdenour Hellas
abdenour.hellas@uqat.ca | +1 418 350 5686
https://abdenour-hellas.online
https://www.linkedin.com/in/abdenour-hellas/
"""

doc.add_paragraph(message)

doc.add_paragraph()

# Liste des recruteurs
doc.add_heading('3. RECRUTEURS ET CONTACTS À CIBLER PAR ENTREPRISE', level=1)

recruiters = [
    ("SNC-Lavalin", "Recruteur / Talent Acquisition Specialist", "Recherchez 'SNC-Lavalin recruiter' sur LinkedIn"),
    ("WSP Canada", "Talent Acquisition - Engineering", "Recherchez 'WSP recruiter engineering'"),
    ("Hydro-Québec", "Recrutement Ingénieurs", "Page carrières Hydro-Québec + LinkedIn"),
    ("Stantec", "Recruiter - Infrastructure", "Recherchez 'Stantec recruiter' sur LinkedIn"),
    ("Rockwell Automation", "Talent Acquisition", "Recherchez 'Rockwell Automation recruiter'"),
    ("BBA", "Recruteur Ingénierie", "Recherchez 'BBA consultant recruiter'"),
    ("Manitoba Hydro", "HR Recruitment", "Page carrières Manitoba Hydro"),
    ("EDF Renewables", "HR Manager", "Recherchez 'EDF Renewables recruiter Canada'"),
    ("Rio Tinto", "Talent Acquisition", "Recherchez 'Rio Tinto recruiter Canada'"),
    ("ABB", "Recruiter - Electrification", "Recherchez 'ABB recruiter electrification'"),
    ("Ontario Power Generation", "Recruitment Specialist", "Page carrières OPG"),
    ("Schneider Electric", "Talent Acquisition", "Recherchez 'Schneider Electric recruiter'"),
    ("BC Hydro", "Recruitment", "Page carrières BC Hydro"),
    ("Emerson", "Talent Acquisition", "Recherchez 'Emerson recruiter Canada'"),
    ("Bombardier", "Recrutement Ingénieurs", "Recherchez 'Bombardier recruiter'"),
    ("Thales Canada", "Recruitment", "Recherchez 'Thales recruiter Canada'"),
    ("Hatch", "Talent Acquisition", "Recherchez 'Hatch recruiter'"),
    ("Pratt & Whitney Canada", "Recruitment", "Recherchez 'PWC recruiter Longueuil'"),
    ("Bruce Power", "Recruitment", "Page carrières Bruce Power"),
    ("Siemens Mobility", "Talent Acquisition", "Recherchez 'Siemens recruiter Montréal'"),
]

for i, (company, role, how) in enumerate(recruiters, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. {company}").bold = True
    p.add_run(f"\n   Rôle cible : {role}")
    p.add_run(f"\n   Comment trouver : {how}")

doc.add_paragraph()

# Conseils
doc.add_heading('4. CONSEILS POUR CONTACTER LES RECRUTEURS', level=1)

doc.add_paragraph('Ne copiez-collez pas le même message à tout le monde. Personnalisez avec le prénom du recruteur.', style='List Number')
doc.add_paragraph('Mentionnez une raison spécifique de contacter CETTE entreprise (projet, valeur, culture).', style='List Number')
doc.add_paragraph('Envoyez les messages mardi-jeudi entre 9h et 11h (meilleur taux d\'ouverture).', style='List Number')
doc.add_paragraph('Relancez poliment après 1 semaine si pas de réponse.', style='List Number')
doc.add_paragraph('Acceptez les demandes de connexion des recruteurs même sans échange préalable.', style='List Number')
doc.add_paragraph('Partagez du contenu pertinent sur LinkedIn (articles sur le génie électrique, l\'automatisation).', style='List Number')

doc.add_paragraph()

# Relance
doc.add_heading('5. MESSAGE DE RELANCE (après 1 semaine)', level=1)

relance = """Bonjour [Prénom],

Je me permis de relancer mon message du [date] concernant les opportunités en génie électrique chez [Entreprise].

Je reste très intéressé par votre équipe et je serais ravi de pouvoir échanger brièvement avec vous, même par téléphone (15 minutes suffisent).

Je vous remercie pour votre temps et reste à votre disposition.

Cordialement,
Abdenour Hellas
"""

doc.add_paragraph(relance)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Bonne chance !').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('CONTACTS_LINKEDIN_RECRUTEURS.docx')
print("✓ CONTACTS_LINKEDIN_RECRUTEURS.docx créé")
