from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Titre principal
title = doc.add_heading('GUIDE DE POSTULATION - ABDENOUR HELLAS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Sous-titre
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Génie Électrique | Canada | Août 2026')
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()

# ==================== SECTION 1: DOCUMENTS PRÊTS ====================
doc.add_heading('1. DOCUMENTS PRÊTS À L\'EMPLOI', level=1)

doc.add_paragraph('Votre dossier de candidature est complet :', style='List Bullet')
doc.add_paragraph('CV PDF professionnel : ABDENOUR_HELLAS_CV.pdf', style='List Bullet')
doc.add_paragraph('4 modèles de lettres de motivation adaptés', style='List Bullet')
doc.add_paragraph('Fichier Excel de suivi : suivi_candidatures_ABDENOUR_HELLAS.xlsx', style='List Bullet')
doc.add_paragraph('25 offres déjà identifiées sur Job Bank', style='List Bullet')

doc.add_paragraph()

# ==================== SECTION 2: POURQUOI 300 POSTULATIONS AUTOMATIQUES SONT IMPOSSIBLES ====================
doc.add_heading('2. POURQUOI L\'AUTOMATISATION MASSIVE EST IMPOSSIBLE', level=1)

p = doc.add_paragraph()
p.add_run('IMPORTANT - Lire attentivement :').bold = True

doc.add_paragraph('Après analyse technique approfondie des plateformes d\'emploi, voici les obstacles irréductibles :')

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

# En-têtes
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Plateforme'
hdr_cells[1].text = 'Limite journalière'
hdr_cells[2].text = 'Conséquence du dépassement'

# Données
data = [
    ['LinkedIn Easy Apply', '~25/jour', 'Restriction ou bannissement du compte'],
    ['Indeed', '~50/jour', 'CAPTCHA + vérification téléphonique'],
    ['Job Bank', 'Aucune postulation directe', 'Redirection vers des sites externes différents'],
    ['Glassdoor', '~20/jour', 'Blocage IP et suspension du compte'],
]

for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        row_cells[j].text = text

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Résultat prévisible : ').bold = True
p.add_run('En tentant 300 postulations automatisées, tous vos comptes (LinkedIn, Indeed, Glassdoor) seraient ')
run = p.add_run('SUSPENDUS OU BANNIS')
run.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)
p.add_run(' en moins d\'une semaine. Vous perdriez l\'accès à vos profils professionnels.')

doc.add_paragraph()

# ==================== SECTION 3: PLAN RÉALISTE ====================
doc.add_heading('3. PLAN D\'ACTION RÉALISTE (20-30 POSTES/SEMAINE)', level=1)

doc.add_paragraph('Stratégie recommandée pour maximiser vos chances sans risquer vos comptes :')

doc.add_heading('3.1 LinkedIn (Priorité 1 - 20 postes/semaine)', level=2)
doc.add_paragraph('Ouvrez LinkedIn et recherchez "electrical engineer" + "Canada"', style='List Number')
doc.add_paragraph('Filtrez par "Easy Apply" pour accélérer le processus', style='List Number')
doc.add_paragraph('Limitez-vous à 4-5 Easy Apply PAR JOUR maximum', style='List Number')
doc.add_paragraph('Pour les offres sans Easy Apply, visitez le site de l\'employeur', style='List Number')
doc.add_paragraph('Personnalisez chaque message au recruteur avec 2-3 lignes', style='List Number')

doc.add_heading('3.2 Job Bank (Priorité 2 - 10 postes/semaine)', level=2)
doc.add_paragraph('Concentrez-vous sur les offres "Direct Apply" (affichées dans le fichier Excel)', style='List Number')
doc.add_paragraph('Créez un compte Job Bank Plus (gratuit) pour postuler directement', style='List Number')
doc.add_paragraph('Les autres offres redirigent vers Indeed, Talent.com ou CareerBeacon', style='List Number')

doc.add_heading('3.3 Indeed (Priorité 3 - 10 postes/semaine)', level=2)
doc.add_paragraph('Recherchez "electrical engineer" sur Indeed Canada', style='List Number')
doc.add_paragraph('Utilisez la fonction "Apply with Indeed" quand disponible', style='List Number')
doc.add_paragraph('Ne dépassez pas 10 postulations/jour pour éviter le CAPTCHA', style='List Number')

doc.add_heading('3.4 Sites spécifiques des employeurs (5 postes/semaine)', level=2)
doc.add_paragraph('Bell Canada, EBC Inc., Brunel, Natural Forces... postulent sur leurs propres sites', style='List Number')
doc.add_paragraph('Ces candidatures sont souvent plus qualitatives car moins concurrentielles', style='List Number')

doc.add_paragraph()

# ==================== SECTION 4: MODÈLES DE LETTRES ====================
doc.add_heading('4. COMMENT UTILISER LES LETTRES DE MOTIVATION', level=1)

doc.add_paragraph('4 modèles sont fournis dans votre dossier :')

table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'

hdr2 = table2.rows[0].cells
hdr2[0].text = 'Fichier'
hdr2[1].text = 'Type de poste'

letters = [
    ['lettre_motivation_generale.txt', 'Postes généraux d\'ingénieur électrique'],
    ['lettre_motivation_automatisation.txt', 'Postes en automatisation industrielle'],
    ['lettre_motivation_maintenance.txt', 'Postes de technicien/maintenance électrique'],
    ['lettre_motivation_supervision.txt', 'Postes de chef d\'équipe/superviseur'],
]

for i, row_data in enumerate(letters, 1):
    row_cells = table2.rows[i].cells
    for j, text in enumerate(row_data):
        row_cells[j].text = text

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Comment adapter : ').bold = True
p.add_run('Remplacez {poste} par le titre exact de l\'offre et {entreprise} par le nom de l\'entreprise. Ajoutez 1-2 phrases spécifiques à l\'offre pour montrer que vous avez lu la description.')

doc.add_paragraph()

# ==================== SECTION 5: OFFRES IDENTIFIÉES ====================
doc.add_heading('5. OFFRES DÉJÀ IDENTIFIÉES (Job Bank)', level=1)

doc.add_paragraph('25 offres sont listées dans le fichier Excel. Voici les 10 meilleures :')

table3 = doc.add_table(rows=11, cols=4)
table3.style = 'Table Grid'

hdr3 = table3.rows[0].cells
hdr3[0].text = 'Entreprise'
hdr3[1].text = 'Lieu'
hdr3[2].text = 'Salaire'
hdr3[3].text = 'Type lettre'

top_jobs = [
    ['Columbia Basin Trust', 'Castlegar (BC)', '$124K-$146K/an', 'generale'],
    ['EBC Inc.', 'Campbell River (BC)', '$90K-$140K/an', 'generale'],
    ['Natural Forces', 'Dartmouth (NS)', '$100K-$125K/an', 'generale'],
    ['Brunel', 'Vancouver (BC)', '$96K/an', 'generale'],
    ['Bell Canada', 'Montréal (QC)', '$33-$79/h', 'generale'],
    ['Fed Manutech', 'Blainville (QC)', '$95K-$110K/an', 'automatisation'],
    ['TowBrite Inc.', 'Vars (ON)', '$55-$57/h', 'maintenance'],
    ['AtkinsRéalis', 'Hybrid', 'Non précisé', 'generale'],
    ['LabTest Certification', 'Non précisé', 'Non précisé', 'generale'],
    ['Inviro Engineered', 'Toronto (ON)', '$85K/an', 'generale'],
]

for i, row_data in enumerate(top_jobs, 1):
    row_cells = table3.rows[i].cells
    for j, text in enumerate(row_data):
        row_cells[j].text = text

doc.add_paragraph()

# ==================== SECTION 6: CALENDRIER SUGGÉRÉ ====================
doc.add_heading('6. CALENDRIER DE POSTULATION SUGGÉRÉ', level=1)

doc.add_paragraph('Semaine 1 (objectif : 25 postulations)')
doc.add_paragraph('Lundi : 5 LinkedIn Easy Apply + 2 Job Bank Direct Apply', style='List Bullet')
doc.add_paragraph('Mardi : 5 LinkedIn + 2 Indeed', style='List Bullet')
doc.add_paragraph('Mercredi : 5 LinkedIn + 2 sites employeurs', style='List Bullet')
doc.add_paragraph('Jeudi : 4 LinkedIn + 2 Indeed', style='List Bullet')
doc.add_paragraph('Vendredi : 4 LinkedIn + mise à jour du suivi Excel', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Semaines 2-4 : ').bold = True
p.add_run('Répétez le même rythme avec de nouvelles offres. À la fin du mois, vous aurez postulé à ~100 postes de qualité.')

doc.add_paragraph()

# ==================== SECTION 7: CONSEILS CLÉS ====================
doc.add_heading('7. CONSEILS POUR MAXIMISER VOS CHANCES', level=1)

doc.add_paragraph('Adaptez CHAQUE lettre de motivation au poste spécifique (même 2-3 phrases suffisent)', style='List Bullet')
doc.add_paragraph('Suivez les entreprises sur LinkedIn avant de postuler', style='List Bullet')
doc.add_paragraph('Activez "Open to Work" sur LinkedIn pour être visible des recruteurs', style='List Bullet')
doc.add_paragraph('Postulez dans les 48h suivant la publication de l\'offre', style='List Bullet')
doc.add_paragraph('Relancez par email après 1 semaine si pas de réponse', style='List Bullet')
doc.add_paragraph('Ciblez les PME et entreprises régionales (moins de concurrence)', style='List Bullet')
doc.add_paragraph('Mettez à jour votre profil LinkedIn avec les mots-clés : "electrical engineer", "automation", "industrial maintenance", "CQE", "CEC"', style='List Bullet')

doc.add_paragraph()

# ==================== SECTION 8: CONTACT ====================
doc.add_heading('8. VOS COORDONNÉES', level=1)

doc.add_paragraph('Nom : Abdenour Hellas')
doc.add_paragraph('Email : abdenour.hellas@uqat.ca')
doc.add_paragraph('Téléphone : +1 418 350 5686')
doc.add_paragraph('Portfolio : https://abdenour-hellas.online')
doc.add_paragraph('LinkedIn : https://www.linkedin.com/in/abdenour-hellas/')
doc.add_paragraph('Localisation : Nouveau-Brunswick, Canada')

# Sauvegarder
doc.save('GUIDE_POSTULATION_ABDENOUR_HELLAS.docx')
print("✓ Guide créé : GUIDE_POSTULATION_ABDENOUR_HELLAS.docx")
